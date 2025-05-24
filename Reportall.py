import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO

st.set_page_config(page_title="Báo cáo nợ quá hạn", layout="wide")
st.title("\U0001F4B8 Báo cáo nợ quá hạn")

def save_bar_chart(data, x_col, y_col, title):
    fig, ax = plt.subplots()
    ax.bar(data[x_col], data[y_col])
    ax.set_title(title)
    ax.set_ylabel(y_col)
    ax.tick_params(axis='x', rotation=45)
    fig.tight_layout()
    buffer = BytesIO()
    fig.savefig(buffer, format='png')
    plt.close(fig)
    buffer.seek(0)
    return buffer

def save_overall_chart(data, x_col, y_col, title):
    fig, ax = plt.subplots(figsize=(10, 5))
    bars = ax.bar(data[x_col], data[y_col])
    ax.set_title(title)
    ax.set_ylabel(y_col)
    ax.tick_params(axis='x', rotation=45)
    for bar in bars:
        height = bar.get_height()
        ax.annotate(f"{height:,.0f}", xy=(bar.get_x() + bar.get_width() / 2, height),
                    xytext=(0, 3), textcoords="offset points", ha='center', va='bottom')
    fig.tight_layout()
    buffer = BytesIO()
    fig.savefig(buffer, format='png')
    plt.close(fig)
    buffer.seek(0)
    return buffer

def set_table_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._element.get_or_add_tblPr()
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def export_word_report(tong_quan, full_df, top3, bot3, charts, nhan_xet, filename):
    doc = Document()
    doc.add_heading("BÁO CÁO ĐÁNH GIÁ", 0)
    doc.add_heading("Tổng quan", level=1)
    for k, v in tong_quan.items():
        doc.add_paragraph(f"{k}: {v}")
    doc.add_heading("Nhận xét", level=1)
    doc.add_paragraph(nhan_xet)

    def add_table(title, df):
        doc.add_heading(title, level=2)
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                col_name = df.columns[i]
                if col_name.lower() == "số tiền":
                    try:
                        row_cells[i].text = f"{float(val):,.0f}".replace(",", ".")
                    except Exception:
                        row_cells[i].text = str(val)
                else:
                    row_cells[i].text = str(val)
        set_table_border(table)

    add_table("Toàn bộ dữ liệu", full_df)
    add_table("Top 3 số tiền cao nhất", top3)
    add_table("Top 3 số tiền thấp nhất", bot3)

    doc.add_heading("Biểu đồ minh họa", level=1)
    for chart in charts:
        doc.add_picture(chart, width=Inches(5.5))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

uploaded_cut = st.file_uploader("\U0001F4C4 Tải lên file Excel công tác cắt điện", type=["xlsx"], key="catdien")
if uploaded_cut:
    df_cut = pd.read_excel(uploaded_cut)
    df_cut.columns = df_cut.columns.str.strip()
    if "Điện lực" not in df_cut.columns:
        st.error("❌ Không tìm thấy cột 'Điện lực' trong file. Hãy kiểm tra lại.")
        st.stop()
    df_cut = df_cut.dropna(subset=["Điện lực"])
    df_cut = df_cut[df_cut["Điện lực"].str.upper() != "TỔNG"]
    df_cut["Khách hàng nợ quá hạn chưa cắt điện"] = df_cut["Khách hàng nợ quá hạn chưa cắt điện"].astype(int)
    df_cut["Số tiền"] = df_cut["Số tiền"].astype(float)
    tong_kh = df_cut["Khách hàng nợ quá hạn chưa cắt điện"].sum()
    tong_tien = df_cut["Số tiền"].sum()
    tong_kh_str = f"{tong_kh:,}".replace(",", ".")
    tong_tien_str = f"{tong_tien:,.0f}".replace(",", ".")
    nhan_xet = f"Hiện tại còn {tong_kh_str} khách hàng chưa bị cắt điện với tổng số tiền nợ {tong_tien_str} đ. Cần rà soát các đơn vị có số tiền lớn để ưu tiên xử lý."

    top_tien = df_cut.sort_values(by="Số tiền", ascending=False).head(3)
    bot_tien = df_cut.sort_values(by="Số tiền", ascending=True).head(3)
    chart_top_tien = save_bar_chart(top_tien, "Điện lực", "Số tiền", "Top 3 Số tiền nợ cao")
    chart_bot_tien = save_bar_chart(bot_tien, "Điện lực", "Số tiền", "Bottom 3 Số tiền nợ thấp")
    chart_all_tien = save_overall_chart(df_cut, "Điện lực", "Số tiền", "Tổng hợp Số tiền nợ")

    st.metric("Tổng KH chưa cắt", tong_kh_str)
    st.metric("Tổng số tiền", tong_tien_str + " đ")
    st.info(nhan_xet)
    st.subheader("\U0001F4C8 Biểu đồ")
    st.image(chart_top_tien)
    st.image(chart_bot_tien)
    st.image(chart_all_tien)

    if st.button("\U0001F4C4 Xuất báo cáo Word", key="tab_word"):
        tong_quan = {
            "Tổng KH chưa cắt": tong_kh_str,
            "Tổng số tiền": tong_tien_str + " đ"
        }
        word_file = export_word_report(
            tong_quan,
            df_cut[["Điện lực", "Khách hàng nợ quá hạn chưa cắt điện", "Số tiền"]],
            top_tien,
            bot_tien,
            [chart_top_tien, chart_bot_tien, chart_all_tien],
            nhan_xet,
            filename="Bao_cao_NoQuaHan.docx"
        )
        st.download_button("\U0001F4E5 Tải báo cáo Word", data=word_file, file_name="Bao_cao_NoQuaHan.docx")
