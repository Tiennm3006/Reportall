import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from datetime import datetime
from docx import Document
from io import BytesIO
import openpyxl
from openpyxl import Workbook
from openpyxl.chart import BarChart, Reference, Series
from openpyxl.utils.dataframe import dataframe_to_rows
from chromadb import PersistentClient
import base64

st.set_page_config(page_title="Phân tích Điện lực", layout="wide")
st.title("📊 Đánh giá công tác kiểm tra và thu tiền điện")

# Tabs chia hai phần
tab1, tab2 = st.tabs(["📋 Hệ thống đo đếm", "💵 Thu tiền điện"])

# ---------- KHỞI TẠO CHROMADB ---------- #
client = PersistentClient(path="./chroma_storage")
collection = client.get_or_create_collection(name="baocao_files")

# ---------- TAB 1: HỆ THỐNG ĐO ĐẾM ---------- #
with tab1:
    uploaded_file = st.file_uploader("Tải lên file Excel chứa sheet 'Tong hop luy ke'", type=["xlsx"], key="kiemtra")

    if uploaded_file:
        df = pd.read_excel(uploaded_file, sheet_name="Tong hop luy ke", header=None)

        # Làm sạch dữ liệu
        df_cleaned = df.iloc[4:].copy()
        df_cleaned.columns = [
            "STT", "Điện lực", "1P_GT", "1P_TT", "3P_GT", "3P_TT",
            "TU", "TI", "Tổng công tơ", "Kế hoạch", "Tỷ lệ"
        ]
        df_cleaned = df_cleaned[df_cleaned["Điện lực"].notna()]
        cols_to_numeric = ["1P_GT", "1P_TT", "3P_GT", "3P_TT", "TU", "TI", "Tổng công tơ", "Kế hoạch", "Tỷ lệ"]
        df_cleaned[cols_to_numeric] = df_cleaned[cols_to_numeric].apply(pd.to_numeric, errors='coerce')
        df_cleaned["Tỷ lệ"] = df_cleaned["Tỷ lệ"] * 100

        # Tổng quan và dự báo
        total_current = df_cleaned["Tổng công tơ"].sum()
        total_plan = df_cleaned["Kế hoạch"].sum()
        current_date = datetime.now()
        days_passed = (current_date - datetime(2025, 1, 1)).days
        days_total = (datetime(2025, 9, 30) - datetime(2025, 1, 1)).days
        avg_per_day = total_current / days_passed
        forecast_total = avg_per_day * days_total
        forecast_ratio = forecast_total / total_plan

        df_sorted = df_cleaned.sort_values(by="Tỷ lệ", ascending=False)
        top_3 = df_sorted.head(3)
        bottom_3 = df_sorted.tail(3)

        st.subheader("Tổng quan và dự báo")
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Tổng đã thực hiện", f"{total_current:,}")
            st.metric("Kế hoạch", f"{total_plan:,}")
            st.metric("Tốc độ TB/ngày", f"{avg_per_day:.2f}")
        with col2:
            st.metric("Dự báo đến 30/09/2025", f"{int(forecast_total):,}")
            st.metric("Tỷ lệ dự báo", f"{forecast_ratio*100:.2f}%")

        fig, ax = plt.subplots(figsize=(10, 5))
        bars = ax.bar(df_sorted["Điện lực"], df_sorted["Tỷ lệ"])
        ax.set_ylabel("Tỷ lệ hoàn thành (%)")
        ax.set_title("Tỷ lệ hoàn thành kế hoạch theo Điện lực")
        ax.tick_params(axis='x', rotation=45)
        for bar in bars:
            height = bar.get_height()
            ax.annotate(f"{height:.1f}%", xy=(bar.get_x() + bar.get_width() / 2, height),
                        xytext=(0, 3), textcoords="offset points", ha='center', va='bottom')
        st.pyplot(fig)

        def generate_docx():
            doc = Document()
            doc.add_heading('BÁO CÁO HỆ THỐNG ĐO ĐẾM', 0)
            doc.add_paragraph(f"Tổng công tơ: {total_current:,}, Kế hoạch: {total_plan:,}, Tỷ lệ: {forecast_ratio*100:.2f}%")
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer

        word_file = generate_docx()
        st.download_button("📄 Tải báo cáo Word", data=word_file, file_name="Bao_cao_kiemtra.docx")

        # Lưu file vào ChromaDB dạng base64
        if st.button("📥 Lưu báo cáo Word vào CSDL nội bộ"):
            encoded = base64.b64encode(word_file.read()).decode("utf-8")
            collection.upsert(
                documents=["Báo cáo kiểm tra ngày " + current_date.strftime('%d/%m/%Y')],
                metadatas=[{"ngay": current_date.strftime('%Y-%m-%d')}],
                ids=[f"baocao_file_{current_date.strftime('%Y%m%d')}"],
            )
            collection.add(
                documents=[encoded],
                metadatas=[{"ngay": current_date.strftime('%Y-%m-%d')}],
                ids=[f"baocao_filedata_{current_date.strftime('%Y%m%d')}"],
            )
            st.success("Đã lưu báo cáo Word vào ChromaDB")

        # Truy vấn lại
        with st.expander("📂 Xem lại các file báo cáo đã lưu"):
            results = collection.get()
            for doc_id, meta in zip(results['ids'], results['metadatas']):
                if doc_id.startswith("baocao_filedata"):
                    ngay = meta['ngay']
                    raw = collection.get(ids=[doc_id])["documents"][0]
                    file_bytes = BytesIO(base64.b64decode(raw))
                    st.download_button(f"📥 Tải báo cáo {ngay}", file_bytes, file_name=f"Bao_cao_{ngay}.docx")

# ---------- TAB 2: THU TIỀN ĐIỆN ---------- #
with tab2:
    st.subheader("Đánh giá công tác thu tiền điện")
    uploaded_payment = st.file_uploader("Tải lên file Excel dữ liệu thu tiền điện", type=["xlsx"], key="thu")

    if uploaded_payment:
        df_pay = pd.read_excel(uploaded_payment)

        st.write("### Dữ liệu gốc")
        st.dataframe(df_pay)

        if {"Điện lực", "Tổng hóa đơn", "Đã thu"}.issubset(df_pay.columns):
            df_pay["Tỷ lệ thu"] = df_pay["Đã thu"] / df_pay["Tổng hóa đơn"] * 100

            st.write("### Tổng hợp")
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Tổng hóa đơn", f"{df_pay['Tổng hóa đơn'].sum():,.0f}")
            with col2:
                st.metric("Tổng đã thu", f"{df_pay['Đã thu'].sum():,.0f}")
            with col3:
                st.metric("Tỷ lệ thu bình quân", f"{df_pay['Tỷ lệ thu'].mean():.2f}%")

            st.write("### Biểu đồ tỷ lệ thu tiền điện")
            fig3, ax3 = plt.subplots()
            ax3.bar(df_pay["Điện lực"], df_pay["Tỷ lệ thu"], color='orange')
            ax3.set_ylabel("Tỷ lệ thu (%)")
            ax3.set_title("Tỷ lệ thu tiền điện theo Điện lực")
            ax3.tick_params(axis='x', rotation=45)
            for i, val in enumerate(df_pay["Tỷ lệ thu"]):
                ax3.annotate(f"{val:.1f}%", xy=(i, val), xytext=(0, 3), textcoords="offset points", ha='center')
            st.pyplot(fig3)
        else:
            st.error("❌ File Excel cần có các cột: 'Điện lực', 'Tổng hóa đơn', 'Đã thu'")
